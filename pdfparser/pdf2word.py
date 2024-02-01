import os
import pdfplumber
import pandas as pd
from gpt_bot_driver import gpt_bot_driver
import time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
pd.set_option('display.max_colwidth', None)

data_dir = '../data/temp'
fnames = os.listdir(data_dir)

# gpt 챗봇 실행
bot = gpt_bot_driver("haesung.kim")

for fname in fnames:
    if fname.endswith('.pdf'):        
        print('processing', fname)
        # 텍스트 파일을 작성할 경로와 파일명 지정
        input_path = os.path.join(data_dir, fname)
        output_path = os.path.join(data_dir, fname).rstrip('.pdf') +'.docx'

        doc = Document()

        pdf = pdfplumber.open(input_path)            
        for page_index, page in enumerate(pdf.pages):
            # 텍스트 추출
            page_text = page.extract_text()
            table = page_text.maketrans('\n',' ') # page \n 제거
            page_text_processed = page_text.translate(table)
            
            # 테이블 추출
            tables = page.extract_tables(
                table_settings = {
                    "vertical_strategy": "lines", 
                    "horizontal_strategy": "lines",
                }
            )
            
            # Raw Text 출력
            doc.add_heading(f'Page {page_index+1}', level=0)
            doc.add_heading(f'Raw Text', level=1)
            doc.add_paragraph(page_text)
            doc.add_page_break()
            
            # 영어 번역 출력
            doc.add_heading(f'Page {page_index+1}', level=0)
            doc.add_heading(f'English Text', level=1)
            bot.delete_conversation_history() # 대화 이력 삭제
            time.sleep(3)
            question = ' Please translate the previous text into English and respond only in English.'
            answer1 = bot.get_gpt_answer(page_text_processed + ' ' + question) # 질의 진행
            doc.add_paragraph(answer1)
            doc.add_page_break()
            
            # 한국어 번역 출력
            doc.add_heading(f'Page {page_index+1}', level=0)
            doc.add_heading(f'Korean Text', level=1)
            bot.delete_conversation_history() # 대화 이력 삭제
            time.sleep(3)
            question = ' 앞의 글을 한국어로 번역해줘, 한국어로만 답변해줘'
            answer2 = bot.get_gpt_answer(page_text_processed + ' ' + question) # 질의 진행
            doc.add_paragraph(answer2)
            doc.add_page_break()
            
            # answer1 \n 제거
            table = answer1.maketrans('\n',' ')
            answer1_processed = answer1.translate(table)
            
            tables = page.extract_tables(
                table_settings = {
                    "vertical_strategy": "lines", 
                    "horizontal_strategy": "lines",
                }
            )

            # 영어 요약 출력
            doc.add_heading(f'Page {page_index+1}', level=0)
            doc.add_heading(f'Summarization', level=1)
            bot.delete_conversation_history() # 대화 이력 삭제
            time.sleep(3)
            question = ' Please summarize the previous text and respond only in English.'
            answer3 = bot.get_gpt_answer(answer1_processed + ' ' + question) # 질의 진행
            doc.add_paragraph(answer3)
            doc.add_page_break()
            
            # 영어 키워드 추출 출력
            doc.add_heading(f'Page {page_index+1}', level=0)
            doc.add_heading(f'Extracted Keyword', level=1)
            bot.delete_conversation_history() # 대화 이력 삭제
            time.sleep(3)
            question = ' Please extract important keywords from the previous text and respond only in English.'
            answer4 = bot.get_gpt_answer(answer1_processed + ' ' + question) # 질의 진행
            doc.add_paragraph(answer4)
            doc.add_page_break()
            
            # Raw Table 출력
            if len(tables):
                doc.add_heading(f'Page {page_index+1}', level=0)
                doc.add_heading(f'Raw Table', level=1)    
                for table_index, table in enumerate(tables):
                    doc.add_heading(f'Table {page_index+1}-{table_index+1}', level=2)
                    
                    # Convert list of lists to dataframe
                    df = pd.DataFrame(table)
                    # add a table to the end and create a reference variable
                    # extra row is so we can add the header row
                    t = doc.add_table(df.shape[0]+1, df.shape[1])

                    for j in range(df.shape[-1]):
                        t.cell(0,j).text = str(df.columns[j])

                    for i in range(df.shape[0]):
                        for j in range(df.shape[-1]):
                            t.cell(i+1,j).text = str(df.values[i,j])
                    # print(f'-----------------------------------------------------English Table-----------------------------------------------------')
                    # print()
                    # print(f'----------------------------------------------------- Korean Table-----------------------------------------------------')
                    # print()
                doc.add_page_break()
        doc.save(output_path)
    